"""Convert README.md to a formatted Word document (.docx)."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

README = Path(__file__).parent / "README.md"
OUTPUT = Path(__file__).parent / "README.docx"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)


def add_code_block(doc, lines):
    """Add a code block as a single paragraph with monospace font and grey background."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # Grey background via highlight
    rPr = run._r.get_or_add_rPr()
    shading = rPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    rPr.append(shading)


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for r, row_data in enumerate(data_rows):
        for c, text in enumerate(row_data):
            if c < cols:
                cell = table.rows[r + 1].cells[c]
                # Handle bold markers
                clean = text.strip()
                cell_para = cell.paragraphs[0]
                # Simple bold handling for **text**
                parts = re.split(r'\*\*(.+?)\*\*', clean)
                for j, part in enumerate(parts):
                    if part:
                        run = cell_para.add_run(part)
                        run.font.size = Pt(9)
                        if j % 2 == 1:  # odd parts are bold
                            run.bold = True

    # Set sensible column widths
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def is_separator_row(line):
    """Check if a markdown table line is a separator (|---|---|)."""
    return bool(re.match(r'^[\s|:-]+$', line.strip().replace('-', '')))


def add_rich_paragraph(doc, text, style=None):
    """Add a paragraph handling inline markdown: **bold**, `code`, [links](url)."""
    p = doc.add_paragraph(style=style)
    # Pattern matches **bold**, `code`, and [text](url)
    pattern = r'(\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))'
    pos = 0
    for m in re.finditer(pattern, text):
        # Add plain text before this match
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.size = Pt(10)
        if m.group(2):  # bold
            run = p.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(10)
        elif m.group(3):  # code
            run = p.add_run(m.group(3))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        elif m.group(4):  # link
            run = p.add_run(m.group(4))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        pos = m.end()
    # Remaining text
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.size = Pt(10)
    return p


def convert():
    md = README.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            h = doc.add_heading(text, level=level)
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing ```
            continue

        # Tables
        if '|' in stripped and not stripped.startswith('>'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                header = parse_table_row(table_lines[0])
                data = []
                for tl in table_lines[1:]:
                    if not is_separator_row(tl):
                        data.append(parse_table_row(tl))
                add_table(doc, header, data)
            continue

        # Blockquotes
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = add_rich_paragraph(doc, quote_text)
            p.paragraph_format.left_indent = Cm(1)
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            # Add a left border via italic styling
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Ordered list items
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_match:
            add_rich_paragraph(doc, f"{ol_match.group(1)}. {ol_match.group(2)}")
            i += 1
            continue

        # Unordered list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            add_rich_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # Indented sub-list items
        indent_match = re.match(r'^(\s{2,})[-*]\s+(.+)$', line)
        if indent_match:
            text = indent_match.group(2)
            p = add_rich_paragraph(doc, text, style='List Bullet')
            indent_level = len(indent_match.group(1)) // 2
            p.paragraph_format.left_indent = Cm(1.0 * indent_level)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, stripped)
        i += 1

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    convert()
